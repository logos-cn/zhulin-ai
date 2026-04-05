from __future__ import annotations

import copy
import hashlib
import json
import logging
import math
import os
import re
import time
from concurrent.futures import Future, ThreadPoolExecutor, TimeoutError as FutureTimeoutError
from dataclasses import dataclass, field
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Iterable, Iterator, Optional
from uuid import uuid4

from sqlalchemy import select
from sqlalchemy.orm import Session

import ai_service
from character_cards import (
    CHARACTER_CARD_TEXT_FIELDS,
    coerce_chapter_number,
    merge_character_card_json,
    normalize_character_timeline_entries,
)
from database import SessionLocal
from models import (
    AIModule,
    Book,
    Character,
    Chapter,
    ChapterNodeType,
    Faction,
    FactionMembership,
    Relation,
    RelationEvent,
    User,
    WorldConflictStrategy,
    WorldExtractionJob,
    WorldExtractionJobStatus,
    WorldExtractionSource,
)
from world_relations import record_relation_event
from world_schema import (
    normalize_faction_status,
    normalize_relation_importance,
    normalize_relation_label,
    normalize_relation_type,
    relation_type_label,
)


WORLD_IMPORT_ROOT = Path(__file__).resolve().parent / "data" / "world_imports"
SUPPORTED_WORLD_IMPORT_EXTENSIONS = {".txt", ".docx", ".pdf"}
DEFAULT_SEGMENT_UNIT_LIMIT = 36000
MIN_SEGMENT_UNIT_LIMIT = 1200
MAX_SEGMENT_UNIT_LIMIT = 60000
MAX_WORLD_EXTRACTION_WORKERS = max(1, int(os.getenv("WORLD_EXTRACTION_WORKER_CAP", "8")))
DEFAULT_IMPORTED_DOCUMENT_WORKERS = max(
    1,
    min(MAX_WORLD_EXTRACTION_WORKERS, int(os.getenv("WORLD_EXTRACTION_MAX_WORKERS", "4"))),
)
DEFAULT_INTERNAL_BOOK_WORKERS = max(
    1,
    min(
        MAX_WORLD_EXTRACTION_WORKERS,
        int(
            os.getenv(
                "WORLD_EXTRACTION_INTERNAL_MAX_WORKERS",
                "2",
            )
        ),
    ),
)
MAX_PARALLEL_EXTRACTION_BUFFER = 8
WORLD_FACT_SUMMARY_MAX_INPUT_UNITS = max(200, int(os.getenv("WORLD_FACT_SUMMARY_MAX_INPUT_UNITS", "12000")))
WORLD_FACT_SUMMARY_CHUNK_TARGET_UNITS = max(
    100,
    int(os.getenv("WORLD_FACT_SUMMARY_CHUNK_TARGET_UNITS", str(max(2000, WORLD_FACT_SUMMARY_MAX_INPUT_UNITS // 2)))),
)
WORLD_FACT_SUMMARY_MAX_PASSES = max(1, int(os.getenv("WORLD_FACT_SUMMARY_MAX_PASSES", "4")))
JOB_HEARTBEAT_INTERVAL_SECONDS = max(5, int(os.getenv("WORLD_EXTRACTION_HEARTBEAT_INTERVAL_SECONDS", "10")))
JOB_HEARTBEAT_STALE_SECONDS = max(
    JOB_HEARTBEAT_INTERVAL_SECONDS * 3,
    int(os.getenv("WORLD_EXTRACTION_HEARTBEAT_STALE_SECONDS", "45")),
)
SERVICE_INSTANCE_ID = os.getenv("WORLD_EXTRACTION_SERVICE_INSTANCE_ID") or uuid4().hex


logger = logging.getLogger("bamboo_ai.world_extraction")

WORLD_EXTRACTION_CHAPTER_STATE_KEY = "world_extraction"
STRICT_JSON_MAX_ATTEMPTS = 3
STRICT_JSON_RETRY_BASE_SECONDS = 0.75
IMPORT_ESTIMATE_INPUT_LOW_FACTOR = 3.8
IMPORT_ESTIMATE_INPUT_HIGH_FACTOR = 7.2
IMPORT_ESTIMATE_SEGMENT_OVERHEAD_LOW = 2400
IMPORT_ESTIMATE_SEGMENT_OVERHEAD_HIGH = 5400
IMPORT_ESTIMATE_OUTPUT_LOW = 800
IMPORT_ESTIMATE_OUTPUT_HIGH = 2400
IMPORT_ESTIMATE_POSTPROCESS_LOW = 700
IMPORT_ESTIMATE_POSTPROCESS_HIGH = 2400
IMPORT_ESTIMATE_WORLD_BIBLE_LOW_MULTIPLIER = 1.12
IMPORT_ESTIMATE_WORLD_BIBLE_HIGH_MULTIPLIER = 1.22


@dataclass
class BookPromptSnapshot:
    title: str
    genre: Optional[str]
    world_bible: Optional[str]


@dataclass
class CharacterPromptSnapshot:
    id: Optional[int]
    name: str
    aliases: list[str] = field(default_factory=list)
    role_label: Optional[str] = None
    description: Optional[str] = None
    traits: list[str] = field(default_factory=list)
    background: Optional[str] = None
    goals: Optional[str] = None
    secrets: Optional[str] = None
    notes: Optional[str] = None
    card_json: dict[str, Any] = field(default_factory=dict)


@dataclass
class ExtractedSegmentPayload:
    segment_label: str
    segment_units: int
    characters: list[dict[str, Any]]
    relations: list[dict[str, Any]]
    factions: list[dict[str, Any]] = field(default_factory=list)
    world_facts: list[str] = field(default_factory=list)


@dataclass
class CharacterTimelineContext:
    chapter_number: Optional[int]
    chapter_label: Optional[str]
    chapter_title: Optional[str]


@dataclass
class ExtractionBlock:
    label: str
    text: str
    chapter_id: Optional[int] = None


@dataclass
class ExtractionSegment:
    label: str
    text: str
    unit_count: int
    chapter_id: Optional[int] = None


class WorldExtractionCancellationRequested(RuntimeError):
    pass


def validate_world_import_source(filename: str) -> str:
    extension = Path(filename or "").suffix.lower()
    if extension == ".doc":
        raise RuntimeError("Legacy .doc is not supported yet. Please resave the file as .docx and upload again.")
    if extension not in SUPPORTED_WORLD_IMPORT_EXTENSIONS:
        raise RuntimeError("Only TXT, DOCX, and PDF files are supported for world extraction.")

    if extension == ".docx":
        try:
            from docx import Document  # noqa: F401
        except ImportError as exc:
            raise RuntimeError("DOCX parsing dependency is missing. Install `python-docx`.") from exc

    if extension == ".pdf":
        try:
            from pypdf import PdfReader  # noqa: F401
        except ImportError as exc:
            raise RuntimeError("PDF parsing dependency is missing. Install `pypdf`.") from exc

    return extension


def normalize_segment_unit_limit(value: Optional[int]) -> int:
    if value is None:
        return DEFAULT_SEGMENT_UNIT_LIMIT
    return max(MIN_SEGMENT_UNIT_LIMIT, min(MAX_SEGMENT_UNIT_LIMIT, int(value)))


def recommended_worker_count(job: WorldExtractionJob) -> int:
    if job.source_type == WorldExtractionSource.IMPORTED_DOCUMENT:
        return DEFAULT_IMPORTED_DOCUMENT_WORKERS
    return DEFAULT_INTERNAL_BOOK_WORKERS


def _effective_worker_count(requested_workers: int, total_segments: int) -> int:
    return max(1, min(max(1, int(requested_workers or 1)), max(1, int(total_segments or 0))))


def job_skip_unchanged_chapters(job: WorldExtractionJob) -> bool:
    options = job.options_json or {}
    if "skip_unchanged_chapters" in options:
        return bool(options.get("skip_unchanged_chapters"))
    return job.source_type == WorldExtractionSource.INTERNAL_BOOK


def _serialize_character_prompt_snapshot(character: Character | CharacterPromptSnapshot) -> CharacterPromptSnapshot:
    raw_card_json = merge_character_card_json(getattr(character, "card_json", None))
    return CharacterPromptSnapshot(
        id=getattr(character, "id", None),
        name=str(getattr(character, "name", "") or "").strip(),
        aliases=list(getattr(character, "aliases", None) or []),
        role_label=getattr(character, "role_label", None),
        description=getattr(character, "description", None),
        traits=list(getattr(character, "traits", None) or []),
        background=getattr(character, "background", None),
        goals=getattr(character, "goals", None),
        secrets=getattr(character, "secrets", None),
        notes=getattr(character, "notes", None),
        card_json=raw_card_json,
    )


def _build_book_prompt_snapshot(book: Book) -> BookPromptSnapshot:
    return BookPromptSnapshot(
        title=book.title,
        genre=book.genre,
        world_bible=book.world_bible,
    )


def _context_aware_segment_unit_limit(
    requested_limit: int,
    *,
    character_config: ai_service.ResolvedAIConfig,
    relation_config: ai_service.ResolvedAIConfig,
) -> tuple[int, Optional[int]]:
    requested = normalize_segment_unit_limit(requested_limit)

    discovered_contexts: list[int] = []
    for config in (character_config, relation_config):
        try:
            metadata = ai_service.get_openai_compatible_model_metadata(
                base_url=config.base_url,
                api_key=config.api_key,
                model_name=config.model_name,
                timeout_seconds=min(config.timeout_seconds, 30),
            )
        except Exception:
            metadata = None
        context_window = metadata.get("context_window") if isinstance(metadata, dict) else None
        if isinstance(context_window, int) and context_window > 0:
            discovered_contexts.append(context_window)

    if not discovered_contexts:
        return requested, None

    context_window = min(discovered_contexts)
    reserved_tokens = max(4000, int(context_window * 0.30))
    usable_tokens = max(4000, context_window - reserved_tokens)
    suggested_units = normalize_segment_unit_limit(int(usable_tokens * 0.40))
    return max(requested, suggested_units), context_window


def _iter_ordered_parallel_extractions(
    segments: Iterator[ExtractionSegment],
    *,
    max_workers: int,
    extractor,
    wait_timeout_seconds: Optional[float] = None,
    on_wait=None,
    should_stop=None,
) -> Iterator[tuple[ExtractionSegment, Optional[ExtractedSegmentPayload], Optional[Exception]]]:
    worker_count = max(1, max_workers)
    buffer_limit = max(worker_count, min(MAX_PARALLEL_EXTRACTION_BUFFER, worker_count * 2))
    executor = ThreadPoolExecutor(max_workers=worker_count, thread_name_prefix="world-extract")
    stop_requested = False
    try:
        pending: list[tuple[ExtractionSegment, Future[ExtractedSegmentPayload]]] = []

        def fill_buffer() -> None:
            while len(pending) < buffer_limit:
                try:
                    next_segment = next(segments)
                except StopIteration:
                    break
                pending.append((next_segment, executor.submit(extractor, next_segment)))

        fill_buffer()
        while pending:
            current_segment, future = pending.pop(0)
            while True:
                if should_stop is not None and should_stop():
                    stop_requested = True
                    future.cancel()
                    for _segment, pending_future in pending:
                        pending_future.cancel()
                    raise WorldExtractionCancellationRequested("World extraction job cancellation requested.")
                try:
                    if wait_timeout_seconds is None or wait_timeout_seconds <= 0:
                        result = future.result()
                    else:
                        result = future.result(timeout=wait_timeout_seconds)
                    yield current_segment, result, None
                    break
                except FutureTimeoutError:
                    if should_stop is not None and should_stop():
                        stop_requested = True
                        future.cancel()
                        for _segment, pending_future in pending:
                            pending_future.cancel()
                        raise WorldExtractionCancellationRequested("World extraction job cancellation requested.")
                    if on_wait is not None:
                        on_wait(current_segment)
                    continue
                except Exception as exc:
                    yield current_segment, None, exc
                    break
            fill_buffer()
    finally:
        executor.shutdown(wait=not stop_requested, cancel_futures=stop_requested)


def ensure_world_import_root() -> Path:
    WORLD_IMPORT_ROOT.mkdir(parents=True, exist_ok=True)
    return WORLD_IMPORT_ROOT


def build_upload_storage_path(job_id: int, original_filename: str) -> Path:
    extension = Path(original_filename or "").suffix.lower()
    safe_stem = "".join(
        char if char.isalnum() or char in {"-", "_"} else "_"
        for char in Path(original_filename or "source").stem
    ).strip("_") or "source"
    job_dir = ensure_world_import_root() / f"job_{job_id}"
    job_dir.mkdir(parents=True, exist_ok=True)
    return job_dir / f"{safe_stem}{extension}"


def _detect_text_encoding(path: Path) -> str:
    with path.open("rb") as handle:
        sample = handle.read(65536)
    for encoding in ("utf-8-sig", "utf-16", "gb18030"):
        try:
            sample.decode(encoding)
            return encoding
        except UnicodeDecodeError:
            continue
    return "utf-8"


def iter_document_blocks(path: Path) -> Iterator[str]:
    extension = validate_world_import_source(path.name)

    if extension == ".txt":
        encoding = _detect_text_encoding(path)
        with path.open("r", encoding=encoding, errors="ignore") as handle:
            buffer: list[str] = []
            for raw_line in handle:
                line = raw_line.strip()
                if not line:
                    if buffer:
                        yield "\n".join(buffer)
                        buffer = []
                    continue
                buffer.append(line)
                if len(buffer) >= 18:
                    yield "\n".join(buffer)
                    buffer = []
            if buffer:
                yield "\n".join(buffer)
        return

    if extension == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX parsing dependency is missing. Install `python-docx`.") from exc
        try:
            document = Document(str(path))
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    yield text
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        yield " | ".join(cells)
        except Exception as exc:
            raise ValueError(f"无法解析上传文档：{path.name}") from exc
        return

    if extension == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("PDF parsing dependency is missing. Install `pypdf`.") from exc
        try:
            reader = PdfReader(str(path))
            for page in reader.pages:
                text = (page.extract_text() or "").strip()
                if text:
                    yield text
        except Exception as exc:
            raise ValueError(f"无法解析上传文档：{path.name}") from exc
        return


def _chapter_sort_key(chapter: Chapter) -> tuple[int, int, int, int]:
    return (
        1 if chapter.sequence_number is None else 0,
        chapter.sequence_number or 0,
        chapter.sort_order,
        chapter.id,
    )


def _chapter_has_extractable_content(chapter: Chapter) -> bool:
    return isinstance(chapter.content, str) and bool(chapter.content.strip())


def _build_internal_chapter_extraction_text(chapter: Chapter) -> str:
    parts = [f"章节标题：{chapter.title}"]
    if chapter.content and chapter.content.strip():
        parts.append(f"章节正文：\n{chapter.content.strip()}")
    return "\n\n".join(part for part in parts if part.strip()).strip()


def iter_internal_book_blocks(
    db: Session,
    book_id: int,
    *,
    chapter_scope: str,
) -> Iterator[ExtractionBlock]:
    chapters = db.execute(select(Chapter).where(Chapter.book_id == book_id)).scalars().all()
    chapters.sort(key=_chapter_sort_key)

    for chapter in chapters:
        if chapter.node_type not in {ChapterNodeType.CHAPTER, ChapterNodeType.SCENE}:
            continue
        if chapter_scope == "with_content" and not _chapter_has_extractable_content(chapter):
            continue

        text = _build_internal_chapter_extraction_text(chapter)
        if text:
            yield ExtractionBlock(
                label=chapter.title,
                text=text,
                chapter_id=chapter.id,
            )


def plan_internal_book_blocks(
    db: Session,
    job: WorldExtractionJob,
) -> tuple[list[ExtractionBlock], dict[str, Any]]:
    chapter_scope = (job.chapter_scope or "with_content").strip() or "with_content"
    skip_unchanged = job_skip_unchanged_chapters(job)
    chapters = db.execute(select(Chapter).where(Chapter.book_id == job.book_id)).scalars().all()
    chapters.sort(key=_chapter_sort_key)

    blocks: list[ExtractionBlock] = []
    stats: dict[str, Any] = {
        "included_chapter_count": 0,
        "skipped_empty_chapter_count": 0,
        "skipped_unchanged_chapter_count": 0,
        "chapter_signatures": {},
    }

    for chapter in chapters:
        if chapter.node_type not in {ChapterNodeType.CHAPTER, ChapterNodeType.SCENE}:
            continue
        if chapter_scope == "with_content" and not _chapter_has_extractable_content(chapter):
            stats["skipped_empty_chapter_count"] += 1
            continue
        if skip_unchanged and not chapter_needs_internal_world_extraction(chapter):
            stats["skipped_unchanged_chapter_count"] += 1
            continue

        text = _build_internal_chapter_extraction_text(chapter)
        if not text:
            stats["skipped_empty_chapter_count"] += 1
            continue

        blocks.append(
            ExtractionBlock(
                label=chapter.title,
                text=text,
                chapter_id=chapter.id,
            )
        )
        stats["included_chapter_count"] += 1
        stats["chapter_signatures"][chapter.id] = _chapter_extraction_signature(chapter)

    return blocks, stats


def _chapter_segment_totals_from_blocks(
    blocks: Iterable[ExtractionBlock],
    *,
    segment_unit_limit: int,
) -> dict[int, int]:
    totals: dict[int, int] = {}
    for block in blocks:
        if block.chapter_id is None:
            continue
        totals[block.chapter_id] = len(_segment_block(block, segment_unit_limit))
    return totals


def _chapter_segment_totals_from_segments(segments: Iterable[ExtractionSegment]) -> dict[int, int]:
    totals: dict[int, int] = {}
    for segment in segments:
        if segment.chapter_id is None:
            continue
        totals[segment.chapter_id] = totals.get(segment.chapter_id, 0) + 1
    return totals


def _segment_block(block: ExtractionBlock, segment_unit_limit: int) -> list[ExtractionSegment]:
    text = block.text.strip()
    if not text:
        return []

    chunks = ai_service.split_text_into_unit_chunks(text, segment_unit_limit)
    if not chunks:
        chunks = [text]

    total = len(chunks)
    segments: list[ExtractionSegment] = []
    for index, chunk in enumerate(chunks, start=1):
        label = block.label if total == 1 else f"{block.label}（片段 {index}/{total}）"
        segments.append(
            ExtractionSegment(
                label=label,
                text=chunk,
                unit_count=ai_service.estimate_text_units(chunk),
                chapter_id=block.chapter_id,
            )
        )
    return segments


def build_segment_plan(
    blocks: Iterable[ExtractionBlock],
    *,
    segment_unit_limit: int,
) -> list[ExtractionSegment]:
    segments: list[ExtractionSegment] = []
    for block in blocks:
        segments.extend(_segment_block(block, segment_unit_limit))
    return segments


def count_segments_from_blocks(
    blocks: Iterable[ExtractionBlock],
    *,
    segment_unit_limit: int,
) -> tuple[int, int]:
    total_units = 0
    total_segments = 0
    for block in blocks:
        total_units += ai_service.estimate_text_units(block.text)
        total_segments += len(_segment_block(block, segment_unit_limit))
    return total_units, total_segments


def iter_segments_from_blocks(
    blocks: Iterable[ExtractionBlock],
    *,
    segment_unit_limit: int,
) -> Iterator[ExtractionSegment]:
    for block in blocks:
        for segment in _segment_block(block, segment_unit_limit):
            yield segment


def _job_failed_segment_labels(job: WorldExtractionJob) -> list[str]:
    errors = (job.result_payload or {}).get("errors") or []
    labels: list[str] = []
    for item in errors:
        if not isinstance(item, dict):
            continue
        label = str(item.get("segment_label") or "").strip()
        if label and label not in labels:
            labels.append(label)
    return labels


def _job_failed_chapter_ids(job: WorldExtractionJob) -> list[int]:
    errors = (job.result_payload or {}).get("errors") or []
    chapter_ids: list[int] = []
    for item in errors:
        if not isinstance(item, dict):
            continue
        value = item.get("chapter_id")
        if isinstance(value, int) and value > 0 and value not in chapter_ids:
            chapter_ids.append(value)
    return chapter_ids


def select_job_retry_segments(
    job: WorldExtractionJob,
    segments: list[ExtractionSegment],
) -> list[ExtractionSegment]:
    options = job.options_json or {}
    if not options.get("retry_failed_only"):
        return segments

    labels = {
        str(item).strip()
        for item in (options.get("failed_segment_labels") or [])
        if str(item).strip()
    }
    if labels:
        matched_by_label = [segment for segment in segments if segment.label in labels]
        if matched_by_label:
            return matched_by_label

    chapter_ids = {
        int(item)
        for item in (options.get("failed_chapter_ids") or [])
        if isinstance(item, int) or (isinstance(item, str) and item.isdigit())
    }
    if chapter_ids:
        return [segment for segment in segments if segment.chapter_id in chapter_ids]

    return []


def _clean_text(value: Any) -> Optional[str]:
    text = str(value or "").strip()
    return text or None


def _contains_cjk_text(value: Any) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", str(value or "")))


def _canonical_relation_text(value: Any) -> str:
    text = str(value or "").strip().lower()
    text = re.sub(r"[_\-]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text


_RELATION_TEXT_TRANSLATIONS = {
    "friend": "朋友",
    "friends": "朋友",
    "ally": "盟友",
    "allies": "盟友",
    "companion": "同伴",
    "companions": "同伴",
    "teammate": "队友",
    "teammates": "队友",
    "classmate": "同学",
    "classmates": "同学",
    "colleague": "同事",
    "colleagues": "同事",
    "roommate": "室友",
    "roommates": "室友",
    "neighbor": "邻居",
    "neighbors": "邻居",
    "relative": "亲属",
    "relatives": "亲属",
    "family": "家人",
    "mentor student": "师徒",
    "master disciple": "师徒",
    "teacher student": "师生",
    "lover": "恋人",
    "lovers": "恋人",
    "couple": "情侣",
    "spouse": "配偶",
    "married": "夫妻",
    "romantic interest": "暧昧对象",
    "enemy": "敌人",
    "enemies": "敌对",
    "rival": "对手",
    "rivals": "对手",
    "boss subordinate": "上下级",
    "leader subordinate": "上下级",
    "employer employee": "雇佣关系",
    "guardian ward": "监护关系",
    "parent child": "亲子",
    "mother son": "母子",
    "mother daughter": "母女",
    "father son": "父子",
    "father daughter": "父女",
    "siblings": "手足",
    "brothers": "兄弟",
    "sisters": "姐妹",
}


def _localize_relation_text(value: Any) -> Optional[str]:
    text = _clean_text(value)
    if not text or _contains_cjk_text(text):
        return text
    return _RELATION_TEXT_TRANSLATIONS.get(_canonical_relation_text(text), text)


def _sanitize_relation_text(value: Any, *, fallback: Optional[str] = None) -> Optional[str]:
    text = _localize_relation_text(value)
    if not text:
        return fallback
    if re.search(r"[A-Za-z]", text):
        return fallback
    return text


def _clean_character_name(value: Any) -> Optional[str]:
    text = _clean_text(value)
    if not text:
        return None
    return re.sub(r"\s+", " ", text)


def _normalized_character_name(value: Any) -> str:
    cleaned = _clean_character_name(value)
    if not cleaned:
        return ""
    return re.sub(r"\s+", "", cleaned).lower()


def _parse_boolean(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.strip().lower() in {"true", "1", "yes", "on"}
    if isinstance(value, int):
        return value != 0
    return bool(value)


def _safe_parse_strength(value: Any) -> Optional[float]:
    if value in (None, ""):
        return None

    try:
        strength = float(value)
    except (TypeError, ValueError):
        logger.warning("invalid_relation_strength value=%r", value)
        return None

    if not math.isfinite(strength):
        logger.warning("non_finite_relation_strength value=%r", value)
        return None

    if strength < 0.0:
        logger.warning("relation_strength_below_range value=%r clamped_to=0.0", value)
        return 0.0
    if strength > 1.0:
        logger.warning("relation_strength_above_range value=%r clamped_to=1.0", value)
        return 1.0
    return strength


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _parse_iso_datetime(value: Any) -> Optional[datetime]:
    text = str(value or "").strip()
    if not text:
        return None
    try:
        parsed = datetime.fromisoformat(text.replace("Z", "+00:00"))
    except ValueError:
        return None
    if parsed.tzinfo is None:
        return parsed.replace(tzinfo=timezone.utc)
    return parsed.astimezone(timezone.utc)


def _job_reference_timestamp(job: WorldExtractionJob) -> Optional[datetime]:
    options = job.options_json or {}
    candidates = [_parse_iso_datetime(options.get("last_heartbeat_at"))]
    if job.status == WorldExtractionJobStatus.RUNNING:
        candidates.extend([job.started_at, job.updated_at, job.created_at])
    for value in candidates:
        if value is None:
            continue
        if value.tzinfo is None:
            return value.replace(tzinfo=timezone.utc)
        return value.astimezone(timezone.utc)
    return None


def _job_has_fresh_heartbeat(job: WorldExtractionJob, *, now: datetime) -> bool:
    reference = _job_reference_timestamp(job)
    if reference is None:
        return False
    return reference >= now - timedelta(seconds=JOB_HEARTBEAT_STALE_SECONDS)


def _update_job_heartbeat(
    job: WorldExtractionJob,
    *,
    db: Session,
    stage: str,
    message: Optional[str] = None,
    commit: bool = False,
) -> None:
    options = dict(job.options_json or {})
    options["service_instance_id"] = SERVICE_INSTANCE_ID
    options["service_pid"] = os.getpid()
    options["last_heartbeat_at"] = _now_iso()
    options["heartbeat_stage"] = stage
    job.options_json = options
    if message is not None:
        job.message = message
    db.add(job)
    if commit:
        db.commit()
        db.refresh(job)


def job_conflict_strategy(job: WorldExtractionJob) -> WorldConflictStrategy:
    if (job.options_json or {}).get("manual_conflict_review"):
        return WorldConflictStrategy.MANUAL_REVIEW
    return job.conflict_strategy


def job_cancel_requested(job: WorldExtractionJob) -> bool:
    return bool((job.options_json or {}).get("cancel_requested"))


def job_is_terminated(job: WorldExtractionJob) -> bool:
    return bool((job.options_json or {}).get("terminated"))


def clear_job_cancel_request(job: WorldExtractionJob) -> None:
    options = dict(job.options_json or {})
    if not options:
        return

    changed = False
    if options.get("cancel_requested"):
        options["cancel_requested"] = False
        changed = True
    if not options.get("terminated") and options.get("termination_requested_at"):
        options.pop("termination_requested_at", None)
        changed = True

    if changed:
        job.options_json = options


def _chapter_world_extraction_state(chapter: Chapter) -> dict[str, Any]:
    extra_data = chapter.extra_data if isinstance(chapter.extra_data, dict) else {}
    state = extra_data.get(WORLD_EXTRACTION_CHAPTER_STATE_KEY)
    return dict(state) if isinstance(state, dict) else {}


def _chapter_extraction_signature(chapter: Chapter) -> str:
    payload = {
        "title": chapter.title or "",
        "content": chapter.content or "",
        "version": int(chapter.version or 0),
        "updated_at": chapter.updated_at.isoformat() if chapter.updated_at else None,
    }
    raw = json.dumps(payload, ensure_ascii=False, sort_keys=True).encode("utf-8")
    return hashlib.sha1(raw).hexdigest()


def chapter_needs_internal_world_extraction(chapter: Chapter) -> bool:
    state = _chapter_world_extraction_state(chapter)
    return state.get("signature") != _chapter_extraction_signature(chapter)


def mark_chapter_internal_world_extracted(
    chapter: Chapter,
    *,
    job_id: int,
    signature: Optional[str] = None,
) -> None:
    extra_data = dict(chapter.extra_data or {})
    extra_data[WORLD_EXTRACTION_CHAPTER_STATE_KEY] = {
        "signature": signature or _chapter_extraction_signature(chapter),
        "job_id": job_id,
        "chapter_version": int(chapter.version or 0),
        "chapter_updated_at": chapter.updated_at.isoformat() if chapter.updated_at else None,
        "extracted_at": _now_iso(),
    }
    chapter.extra_data = extra_data


def mark_job_terminated(
    job: WorldExtractionJob,
    *,
    message: str,
    error_message: str = "已由用户终止。",
) -> None:
    options = dict(job.options_json or {})
    options["cancel_requested"] = False
    options["terminated"] = True
    options.setdefault("termination_reason", "user_cancelled")
    options["terminated_at"] = _now_iso()
    job.options_json = options
    job.status = WorldExtractionJobStatus.FAILED
    job.finished_at = datetime.now(timezone.utc)
    job.error_message = error_message
    job.message = message


def recover_interrupted_world_extraction_jobs(db: Session) -> int:
    now = datetime.now(timezone.utc)
    jobs = db.execute(
        select(WorldExtractionJob).where(
            WorldExtractionJob.status.in_(
                [WorldExtractionJobStatus.PENDING, WorldExtractionJobStatus.RUNNING]
            )
        )
    ).scalars().all()

    recovered = 0
    for job in jobs:
        if _job_has_fresh_heartbeat(job, now=now):
            logger.info(
                "world_extraction_recovery_skipped_fresh_heartbeat job_id=%s status=%s",
                job.id,
                job.status.value if isinstance(job.status, WorldExtractionJobStatus) else job.status,
            )
            continue

        previous_status = job.status
        options = dict(job.options_json or {})
        options["recovered_at"] = _now_iso()
        options["recovery_reason"] = "service_process_interrupted"

        if job_cancel_requested(job):
            mark_job_terminated(
                job,
                message="服务进程中断或重新启动后，提取任务已标记为终止。",
            )
            options.update(job.options_json or {})
            job.options_json = options
            db.add(job)
            recovered += 1
            continue

        job.status = WorldExtractionJobStatus.FAILED
        job.finished_at = datetime.now(timezone.utc)
        job.error_message = "提取任务因服务进程中断或重新启动而中断，可在任务列表中继续提取。"
        job.message = (
            "提取队列在开始前因服务进程中断或重新启动而中断。"
            if previous_status == WorldExtractionJobStatus.PENDING
            else "提取任务在服务进程中断或重新启动时意外停止。"
        )
        job.options_json = options
        db.add(job)
        recovered += 1

    if recovered:
        db.commit()
    return recovered


def _select_text(existing: Any, incoming: Any, strategy: WorldConflictStrategy) -> Optional[str]:
    existing_text = _clean_text(existing)
    incoming_text = _clean_text(incoming)

    if strategy == WorldConflictStrategy.KEEP_EXISTING:
        return existing_text or incoming_text
    if strategy == WorldConflictStrategy.PREFER_IMPORTED:
        return incoming_text or existing_text
    if existing_text and incoming_text:
        return incoming_text if len(incoming_text) >= len(existing_text) else existing_text
    return incoming_text or existing_text


def _select_list(existing: Any, incoming: Any, strategy: WorldConflictStrategy) -> list[str]:
    existing_items = ai_service._string_list(existing)
    incoming_items = ai_service._string_list(incoming)

    if strategy == WorldConflictStrategy.KEEP_EXISTING:
        return existing_items or incoming_items
    if strategy == WorldConflictStrategy.PREFER_IMPORTED:
        return incoming_items or existing_items
    return ai_service._merge_unique_strings(existing_items, incoming_items)


def _character_lookup(characters: Iterable[Character]) -> dict[str, Character]:
    lookup: dict[str, Character] = {}
    for character in characters:
        key = _normalized_character_name(character.name)
        if key and key not in lookup:
            lookup[key] = character
        for alias in character.aliases or []:
            alias_key = _normalized_character_name(alias)
            if alias_key and alias_key not in lookup:
                lookup[alias_key] = character
    return lookup


def _register_character_lookup(lookup: dict[str, Character], character: Character) -> None:
    keys = {_normalized_character_name(character.name)}
    keys.update(_normalized_character_name(alias) for alias in (character.aliases or []))
    for key in keys:
        if key:
            lookup[key] = character


def _segment_timeline_context(
    db: Session,
    *,
    book_id: int,
    segment: ExtractionSegment,
) -> CharacterTimelineContext:
    if segment.chapter_id is not None:
        chapter = db.get(Chapter, segment.chapter_id)
        if chapter is not None and chapter.book_id == book_id:
            chapter_number = chapter.sequence_number or chapter.sort_order or chapter.id
            chapter_label = f"第{chapter_number}章"
            chapter_title = chapter.title or chapter_label
            return CharacterTimelineContext(
                chapter_number=chapter_number,
                chapter_label=chapter_label,
                chapter_title=chapter_title,
            )

    fallback_number = coerce_chapter_number(segment.label)
    if fallback_number is None:
        return CharacterTimelineContext(
            chapter_number=None,
            chapter_label=segment.label or None,
            chapter_title=segment.label or None,
        )

    return CharacterTimelineContext(
        chapter_number=fallback_number,
        chapter_label=f"第{fallback_number}章",
        chapter_title=segment.label or f"第{fallback_number}章",
    )


def _incoming_character_timeline_entries(
    raw_item: dict[str, Any],
    *,
    timeline_context: CharacterTimelineContext,
) -> list[dict[str, Any]]:
    raw_entries = raw_item.get("timeline_entries")
    if isinstance(raw_entries, (list, tuple)):
        stamped_entries: list[dict[str, Any]] = []
        for item in raw_entries:
            if not isinstance(item, dict):
                continue
            entry = dict(item)
            if timeline_context.chapter_number is not None:
                entry["chapter_number"] = timeline_context.chapter_number
            if timeline_context.chapter_label:
                entry["chapter_label"] = timeline_context.chapter_label
            if timeline_context.chapter_title:
                entry["chapter_title"] = timeline_context.chapter_title
            stamped_entries.append(entry)

        normalized_entries = normalize_character_timeline_entries(stamped_entries)
        if normalized_entries:
            return normalized_entries

    event = _clean_text(
        raw_item.get("timeline_event")
        or raw_item.get("current_event")
        or raw_item.get("event")
    )
    location = _clean_text(raw_item.get("current_location") or raw_item.get("location"))
    status = _clean_text(raw_item.get("current_status") or raw_item.get("status"))
    notes = _clean_text(raw_item.get("timeline_notes"))
    if not any((event, location, status, notes)):
        return []

    generated_entry: dict[str, Any] = {}
    if timeline_context.chapter_number is not None:
        generated_entry["chapter_number"] = timeline_context.chapter_number
    if timeline_context.chapter_label:
        generated_entry["chapter_label"] = timeline_context.chapter_label
    if timeline_context.chapter_title:
        generated_entry["chapter_title"] = timeline_context.chapter_title
    if event:
        generated_entry["event"] = event
    if location:
        generated_entry["location"] = location
    if status:
        generated_entry["status"] = status
    if notes:
        generated_entry["notes"] = notes
    return normalize_character_timeline_entries([generated_entry])


def _relation_prompt_candidates(
    existing_characters: list[Character],
    extracted_payloads: list[dict[str, Any]],
    segment_text: str,
    *,
    limit: int = 120,
) -> list[dict[str, Any]]:
    selected: list[dict[str, Any]] = []
    seen: set[str] = set()
    lowered_text = segment_text.lower()

    def add_payload(payload: dict[str, Any]) -> None:
        key = ai_service._normalized_name(payload.get("name"))
        if not key or key in seen:
            return
        seen.add(key)
        selected.append(payload)

    for payload in extracted_payloads:
        add_payload(payload)

    for character in existing_characters:
        names = [character.name, *(character.aliases or [])]
        if any(name and name.lower() in lowered_text for name in names):
            add_payload(ai_service._serialize_character(character))

    if not selected:
        for character in existing_characters[: min(limit, 40)]:
            add_payload(ai_service._serialize_character(character))

    return selected[:limit]


def _world_character_prompt(
    book: Book,
    segment: ExtractionSegment,
    existing_characters: list[Character],
) -> list[dict[str, str]]:
    existing_names = [character.name for character in existing_characters]
    return [
        {
            "role": "system",
            "content": (
                "You extract novel world data. Return strict JSON only. "
                "Use Chinese text when the source is Chinese. "
                "Do not wrap the answer in markdown."
            ),
        },
        {
            "role": "user",
            "content": (
                "Extract characters and world facts from the provided novel segment.\n"
                "Return a JSON object with this shape:\n"
                "{\n"
                '  "characters": [\n'
                "    {\n"
                '      "name": "string",\n'
                '      "aliases": ["string"],\n'
                '      "role_label": "string",\n'
                '      "importance_level": "major | minor",\n'
                '      "description": "string",\n'
                '      "traits": ["string"],\n'
                '      "background": "string",\n'
                '      "goals": "string",\n'
                '      "secrets": "string",\n'
                '      "notes": "string",\n'
                '      "age": "string",\n'
                '      "short_term_goal": "string",\n'
                '      "long_term_goal": "string",\n'
                '      "motivation": "string",\n'
                '      "personality": "string",\n'
                '      "appearance": "string",\n'
                '      "weakness": "string",\n'
                '      "life_statuses": ["alive | dead | serious_injury | minor_injury | disabled"],\n'
                '      "timeline_entries": [\n'
                "        {\n"
                '          "event": "string",\n'
                '          "location": "string",\n'
                '          "status": "string"\n'
                "        }\n"
                "      ]\n"
                "    }\n"
                "  ],\n"
                '  "world_facts": ["string"]\n'
                "}\n\n"
                f"Book title: {book.title}\n"
                f"Genre: {ai_service._safe_text(book.genre, 'unknown')}\n"
                f"Existing characters: {json.dumps(existing_names, ensure_ascii=False)}\n"
                f"Book world bible: {ai_service._safe_text(book.world_bible, 'none')}\n"
                f"Source label: {segment.label}\n"
                f"Segment content:\n{segment.text}\n\n"
                "Merge aliases into existing names when obviously the same person. "
                "Only include characters that are meaningfully present in this segment. "
                "For temporary passersby or low-importance characters, either omit them or mark importance_level as minor. "
                "Only fill age, personality, appearance, weakness, motivation, short_term_goal, long_term_goal, background, secrets, and life_statuses when the segment clearly supports them. "
                "Use short_term_goal for an immediate segment-visible objective, use long_term_goal only for a stable long-range pursuit, and do not guess hidden motivation. "
                "If the segment explicitly shows a stable life state such as alive, dead, serious injury, minor injury, or disability, add it to life_statuses. "
                "Do not guess life_statuses when the segment does not clearly support it. "
                "For each kept character, add at most 1 timeline_entries item describing what they are doing in this segment, where they are, and what state they are in. "
                "timeline_entries should only describe facts visible in this segment, not biography summary or future prediction. "
                "Prefer concrete visible actions and current location/state over vague summaries. "
                "For world_facts, only keep durable setting facts or stable人物/组织/地点事实. "
                "Do not include plot recap, temporary actions, dialogue, or repeated facts. "
                "Return at most 5 world_facts, each as one concise sentence."
            ),
        },
    ]


def _world_relation_prompt(
    book: Book,
    segment: ExtractionSegment,
    characters: list[dict[str, Any]],
) -> list[dict[str, str]]:
    return [
        {
            "role": "system",
            "content": (
                "You extract relationships between known novel characters. "
                "Return strict JSON only, no markdown. "
                "Use stable English keys for relation_type: kinship, affinity, hostility, authority, other. "
                "Use Chinese for label, description, and event_summary when the source text is Chinese. "
                "Keep description to one short sentence, ideally within 18 to 40 Chinese characters."
            ),
        },
        {
            "role": "user",
            "content": (
                "Return a JSON object with this shape:\n"
                "{\n"
                '  "relations": [\n'
                "    {\n"
                '      "source_name": "string",\n'
                '      "target_name": "string",\n'
                '      "relation_type": "kinship | affinity | hostility | authority | other",\n'
                '      "label": "string",\n'
                '      "description": "string",\n'
                '      "strength": 0.0,\n'
                '      "importance_level": "core | major | minor | background",\n'
                '      "event_summary": "string",\n'
                '      "is_bidirectional": false\n'
                "    }\n"
                "  ]\n"
                "}\n\n"
                f"Book title: {book.title}\n"
                f"Source label: {segment.label}\n"
                f"Known characters JSON: {json.dumps(characters, ensure_ascii=False)}\n"
                f"Segment content:\n{segment.text}\n\n"
                "Only use source_name and target_name values from the known characters list. "
                "Skip speculative relationships. "
                "relation_type 必须严格使用指定英文键。"
                "label、description、event_summary 必须优先输出简洁中文，不要输出英文关系词。"
                "description 只保留一条简短关系说明，不要写剧情复述、长摘要或多句分析。"
            ),
        },
    ]


def _merge_character_payload(
    existing: Optional[Character],
    raw_item: dict[str, Any],
    *,
    strategy: WorldConflictStrategy,
    chapter_id: Optional[int],
    timeline_context: Optional[CharacterTimelineContext] = None,
) -> dict[str, Any]:
    cleaned_name = _clean_character_name(raw_item.get("name"))
    if existing is None and not cleaned_name:
        raise ValueError("Character name cannot be empty for new characters.")

    aliases = _select_list(existing.aliases if existing else [], raw_item.get("aliases"), strategy)
    traits = _select_list(existing.traits if existing else [], raw_item.get("traits"), strategy)
    existing_card = existing.card_json or {} if existing else {}
    incoming_card = raw_item if isinstance(raw_item, dict) else {}
    incoming_timeline_entries = _incoming_character_timeline_entries(
        raw_item,
        timeline_context=timeline_context
        or CharacterTimelineContext(chapter_number=None, chapter_label=None, chapter_title=None),
    )

    if strategy == WorldConflictStrategy.KEEP_EXISTING:
        card_json = dict(existing_card or incoming_card)
    else:
        card_json = {**existing_card, **incoming_card}
    card_json = merge_character_card_json(
        card_json,
        life_statuses=card_json.get("life_statuses"),
        timeline_entries=[*(existing_card.get("timeline_entries") or []), *incoming_timeline_entries],
    )

    appearance_points = [
        value
        for value in (
            existing.first_appearance_chapter_id if existing else None,
            existing.last_appearance_chapter_id if existing else None,
            chapter_id,
        )
        if value is not None
    ]
    first_appearance = min(appearance_points) if appearance_points else None
    last_appearance = max(appearance_points) if appearance_points else None

    return {
        "name": existing.name if existing else cleaned_name,
        "aliases": aliases,
        "role_label": _select_text(existing.role_label if existing else None, raw_item.get("role_label"), strategy),
        "description": _select_text(existing.description if existing else None, raw_item.get("description"), strategy),
        "traits": traits,
        "background": _select_text(existing.background if existing else None, raw_item.get("background"), strategy),
        "goals": _select_text(existing.goals if existing else None, raw_item.get("goals"), strategy),
        "secrets": _select_text(existing.secrets if existing else None, raw_item.get("secrets"), strategy),
        "notes": _select_text(existing.notes if existing else None, raw_item.get("notes"), strategy),
        "first_appearance_chapter_id": first_appearance,
        "last_appearance_chapter_id": last_appearance,
        "is_active": True if existing is None else existing.is_active,
        "card_json": card_json,
    }


def _apply_relation_payload(
    relation: Relation,
    payload: dict[str, Any],
    *,
    strategy: WorldConflictStrategy,
) -> None:
    existing_description = ai_service.normalize_relation_description(relation.description)
    incoming_description = ai_service.normalize_relation_description(payload.get("description"))

    if strategy == WorldConflictStrategy.KEEP_EXISTING:
        relation.label = relation.label or payload["label"]
        relation.description = existing_description or incoming_description
        if relation.strength is None:
            relation.strength = payload["strength"]
        relation.importance_level = relation.importance_level or payload["importance_level"]
        relation.is_bidirectional = relation.is_bidirectional or payload["is_bidirectional"]
        return

    if strategy in {WorldConflictStrategy.PREFER_IMPORTED, WorldConflictStrategy.MANUAL_REVIEW}:
        relation.label = payload["label"]
        relation.description = incoming_description
        relation.strength = payload["strength"]
        relation.importance_level = payload["importance_level"]
        relation.is_bidirectional = payload["is_bidirectional"]
        return

    if strategy != WorldConflictStrategy.MERGE:
        raise ValueError(f"Unsupported conflict strategy: {strategy}")

    relation.label = payload["label"] or relation.label
    relation.description = incoming_description or existing_description
    relation.strength = payload["strength"] if payload["strength"] is not None else relation.strength
    relation.importance_level = payload["importance_level"] or relation.importance_level
    relation.is_bidirectional = relation.is_bidirectional or payload["is_bidirectional"]


def _select_existing_relation(
    db: Session,
    *,
    book_id: int,
    source_character_id: int,
    target_character_id: int,
    relation_type: str,
) -> Optional[Relation]:
    matches = db.execute(
        select(Relation)
        .where(
            Relation.book_id == book_id,
            Relation.source_character_id == source_character_id,
            Relation.target_character_id == target_character_id,
            Relation.relation_type == relation_type,
        )
        .order_by(Relation.id.asc())
    ).scalars().all()
    if len(matches) > 1:
        logger.warning(
            "duplicate_relations_detected book_id=%s source_character_id=%s target_character_id=%s relation_type=%s relation_ids=%s",
            book_id,
            source_character_id,
            target_character_id,
            relation_type,
            [item.id for item in matches],
        )
    return matches[0] if matches else None


def _serialize_character_snapshot(character: Character) -> dict[str, Any]:
    normalized_card = merge_character_card_json(character.card_json)
    payload = {
        "id": character.id,
        "name": character.name,
        "aliases": character.aliases or [],
        "role_label": character.role_label,
        "description": character.description,
        "traits": character.traits or [],
        "background": character.background,
        "goals": character.goals,
        "secrets": character.secrets,
        "notes": character.notes,
        "first_appearance_chapter_id": character.first_appearance_chapter_id,
        "last_appearance_chapter_id": character.last_appearance_chapter_id,
        "card_json": normalized_card,
    }
    for key in CHARACTER_CARD_TEXT_FIELDS:
        payload[key] = normalized_card.get(key)
    return payload


def _serialize_relation_snapshot(relation: Relation) -> dict[str, Any]:
    return {
        "id": relation.id,
        "source_character_id": relation.source_character_id,
        "source_character_name": relation.source_character.name if relation.source_character else None,
        "target_character_id": relation.target_character_id,
        "target_character_name": relation.target_character.name if relation.target_character else None,
        "relation_type": normalize_relation_type(relation.relation_type),
        "relation_type_label": relation_type_label(relation.relation_type),
        "label": _sanitize_relation_text(relation.label),
        "description": ai_service.normalize_relation_description(relation.description),
        "strength": relation.strength,
        "importance_level": normalize_relation_importance(relation.importance_level),
        "is_bidirectional": relation.is_bidirectional,
    }


def _relation_conflict_key(source_id: int, target_id: int, relation_type: str) -> str:
    return f"{source_id}:{target_id}:{relation_type.strip().lower()}"


def _add_or_update_conflict(result_conflicts: list[dict[str, Any]], conflict: dict[str, Any]) -> None:
    source_label = conflict.get("source_label")
    for item in result_conflicts:
        if item.get("id") != conflict.get("id"):
            continue
        existing_labels = item.get("source_labels") or []
        if source_label and source_label not in existing_labels:
            existing_labels.append(source_label)
            item["source_labels"] = existing_labels
        item["incoming"] = conflict.get("incoming")
        item["updated_at"] = _now_iso()
        return

    payload = dict(conflict)
    payload["status"] = payload.get("status") or "pending"
    payload["source_labels"] = [source_label] if source_label else []
    payload.pop("source_label", None)
    payload["created_at"] = _now_iso()
    payload["updated_at"] = payload["created_at"]
    result_conflicts.append(payload)


def _recount_conflicts(result_payload: dict[str, Any]) -> None:
    conflicts = result_payload.get("conflicts") or []
    pending = sum(1 for item in conflicts if item.get("status") != "resolved")
    resolved = sum(1 for item in conflicts if item.get("status") == "resolved")
    result_payload.setdefault("totals", {})
    result_payload["totals"]["pending_conflict_count"] = pending
    result_payload["totals"]["resolved_conflict_count"] = resolved


def _call_strict_json_chat(
    config: ai_service.ResolvedAIConfig,
    *,
    messages: list[dict[str, str]],
    expectation: str,
) -> Any:
    last_error: Optional[Exception] = None
    for attempt in range(1, STRICT_JSON_MAX_ATTEMPTS + 1):
        try:
            response = ai_service.call_openai_compatible_chat(
                config,
                messages=messages,
            )
            payload = ai_service._extract_json_block(response["text"])
            return payload
        except ai_service.AIInvocationError as exc:
            last_error = exc
            if attempt >= STRICT_JSON_MAX_ATTEMPTS:
                break
            time.sleep(STRICT_JSON_RETRY_BASE_SECONDS * (2 ** (attempt - 1)))
    raise ai_service.AIInvocationError(f"{expectation}: {last_error}") from last_error


def _world_faction_prompt(
    book: BookPromptSnapshot,
    segment: ExtractionSegment,
    characters: list[dict[str, Any]],
) -> list[dict[str, str]]:
    return [
        {
            "role": "system",
            "content": (
                "You extract durable factions and character memberships from a Chinese novel segment. "
                "Return strict JSON only. "
                "Only keep durable organizations, camps, sects, courts, parties, or forces."
            ),
        },
        {
            "role": "user",
            "content": (
                "Return a JSON object with this shape:\n"
                "{\n"
                '  "factions": [\n'
                "    {\n"
                '      "name": "string",\n'
                '      "description": "string",\n'
                '      "memberships": [\n'
                "        {\n"
                '          "character_name": "string",\n'
                '          "role_label": "string",\n'
                '          "loyalty": 0.0,\n'
                '          "status": "active | former",\n'
                '          "notes": "string"\n'
                "        }\n"
                "      ]\n"
                "    }\n"
                "  ]\n"
                "}\n\n"
                f"Book title: {book.title}\n"
                f"Source label: {segment.label}\n"
                f"Known characters JSON: {json.dumps(characters, ensure_ascii=False)}\n"
                f"Segment content:\n{segment.text}\n\n"
                "Only use character_name values from the known characters list. "
                "Skip temporary one-scene teams. "
                "If no durable faction is evident, return an empty factions list."
            ),
        },
    ]


def extract_segment_world_payload(
    *,
    book_snapshot: BookPromptSnapshot,
    segment: ExtractionSegment,
    existing_character_snapshots: list[CharacterPromptSnapshot],
    character_config: ai_service.ResolvedAIConfig,
    relation_config: ai_service.ResolvedAIConfig,
) -> ExtractedSegmentPayload:
    character_json = _call_strict_json_chat(
        character_config,
        messages=_world_character_prompt(book_snapshot, segment, existing_character_snapshots),
        expectation="Character extraction response must be a JSON object.",
    )
    if not isinstance(character_json, dict):
        raise ai_service.AIInvocationError("Character extraction response must be a JSON object.")

    extracted_characters = character_json.get("characters") or []
    if not isinstance(extracted_characters, list):
        extracted_characters = []
    world_facts = ai_service._string_list(character_json.get("world_facts"))

    relation_candidates = _relation_prompt_candidates(
        existing_character_snapshots,
        [item for item in extracted_characters if isinstance(item, dict)],
        segment.text,
    )
    relation_json = _call_strict_json_chat(
        relation_config,
        messages=_world_relation_prompt(book_snapshot, segment, relation_candidates),
        expectation="Relation extraction response must be a JSON object.",
    )
    if not isinstance(relation_json, dict):
        raise ai_service.AIInvocationError("Relation extraction response must be a JSON object.")

    extracted_relations = relation_json.get("relations") or []
    if not isinstance(extracted_relations, list):
        extracted_relations = []

    faction_json = _call_strict_json_chat(
        relation_config,
        messages=_world_faction_prompt(
            book_snapshot,
            segment,
            relation_candidates,
        ),
        expectation="Faction extraction response must be a JSON object.",
    )
    if not isinstance(faction_json, dict):
        raise ai_service.AIInvocationError("Faction extraction response must be a JSON object.")
    extracted_factions = faction_json.get("factions") or []
    if not isinstance(extracted_factions, list):
        extracted_factions = []

    return ExtractedSegmentPayload(
        segment_label=segment.label,
        segment_units=segment.unit_count,
        characters=[item for item in extracted_characters if isinstance(item, dict)],
        relations=[item for item in extracted_relations if isinstance(item, dict)],
        factions=[item for item in extracted_factions if isinstance(item, dict)],
        world_facts=world_facts,
    )


def apply_segment_world_payload(
    db: Session,
    *,
    book: Book,
    segment: ExtractionSegment,
    extracted_payload: ExtractedSegmentPayload,
    conflict_strategy: WorldConflictStrategy,
    update_world_bible: bool,
) -> dict[str, Any]:
    existing_characters = db.execute(
        select(Character).where(Character.book_id == book.id).order_by(Character.name.asc())
    ).scalars().all()
    character_lookup = _character_lookup(existing_characters)
    timeline_context = _segment_timeline_context(db, book_id=book.id, segment=segment)

    merged_character_payloads: list[dict[str, Any]] = []
    created_character_count = 0
    updated_character_count = 0
    character_conflict_count = 0
    touched_character_ids: set[int] = set()
    conflicts: list[dict[str, Any]] = []

    for raw_item in extracted_payload.characters:
        if not isinstance(raw_item, dict):
            logger.warning("skipping_character_non_mapping payload=%r", raw_item)
            continue

        name = _clean_character_name(raw_item.get("name"))
        if not name:
            logger.warning("skipping_character_empty_name segment=%s payload=%r", segment.label, raw_item)
            continue

        normalized_name = _normalized_character_name(name)
        if not normalized_name:
            logger.warning("skipping_character_non_normalizable_name segment=%s name=%r", segment.label, name)
            continue

        incoming_item = dict(raw_item)
        incoming_item["name"] = name
        existing = character_lookup.get(normalized_name)
        if existing is not None:
            character_conflict_count += 1

        merged_payload = _merge_character_payload(
            existing,
            incoming_item,
            strategy=conflict_strategy,
            chapter_id=segment.chapter_id,
            timeline_context=timeline_context,
        )
        merged_character_payloads.append(merged_payload)

        if existing is not None and conflict_strategy == WorldConflictStrategy.MANUAL_REVIEW:
            conflicts.append(
                {
                    "id": f"character:{existing.id}",
                    "conflict_type": "character",
                    "target_id": existing.id,
                    "title": existing.name,
                    "existing": _serialize_character_snapshot(existing),
                    "incoming": merged_payload,
                    "source_label": segment.label,
                }
            )
            continue

        if existing is None:
            existing = Character(book_id=book.id, name=merged_payload["name"] or name)
            created_character_count += 1
        else:
            updated_character_count += 1

        existing.name = merged_payload["name"] or existing.name
        existing.aliases = merged_payload["aliases"]
        existing.role_label = merged_payload["role_label"]
        existing.description = merged_payload["description"]
        existing.traits = merged_payload["traits"]
        existing.background = merged_payload["background"]
        existing.goals = merged_payload["goals"]
        existing.secrets = merged_payload["secrets"]
        existing.notes = merged_payload["notes"]
        existing.first_appearance_chapter_id = merged_payload["first_appearance_chapter_id"]
        existing.last_appearance_chapter_id = merged_payload["last_appearance_chapter_id"]
        existing.is_active = merged_payload["is_active"]
        existing.card_json = merged_payload["card_json"]
        db.add(existing)
        db.flush()
        touched_character_ids.add(existing.id)
        _register_character_lookup(character_lookup, existing)

    lookup_after_merge = character_lookup
    merged_relation_payloads: list[dict[str, Any]] = []
    created_relation_count = 0
    updated_relation_count = 0
    relation_conflict_count = 0
    touched_relation_ids: set[int] = set()
    created_faction_count = 0
    updated_faction_count = 0
    touched_faction_ids: set[int] = set()
    touched_membership_ids: set[int] = set()

    for raw_item in extracted_payload.relations:
        source_name = ai_service._normalized_name(raw_item.get("source_name"))
        target_name = ai_service._normalized_name(raw_item.get("target_name"))
        relation_type = normalize_relation_type(raw_item.get("relation_type"))
        if not source_name or not target_name or not relation_type or source_name == target_name:
            continue

        source_character = lookup_after_merge.get(source_name)
        target_character = lookup_after_merge.get(target_name)
        if source_character is None or target_character is None:
            continue

        payload = {
            "source_character_id": source_character.id,
            "source_name": source_character.name,
            "target_character_id": target_character.id,
            "target_name": target_character.name,
            "relation_type": relation_type,
            "relation_type_label": relation_type_label(relation_type),
            "label": _sanitize_relation_text(raw_item.get("label")) or relation_type_label(relation_type),
            "description": ai_service.normalize_relation_description(_clean_text(raw_item.get("description"))),
            "strength": _safe_parse_strength(raw_item.get("strength")),
            "importance_level": normalize_relation_importance(raw_item.get("importance_level")),
            "event_summary": ai_service.normalize_relation_description(_clean_text(raw_item.get("event_summary"))),
            "is_bidirectional": _parse_boolean(raw_item.get("is_bidirectional")),
        }
        merged_relation_payloads.append(payload)

        existing_relation = _select_existing_relation(
            db,
            book_id=book.id,
            source_character_id=source_character.id,
            target_character_id=target_character.id,
            relation_type=relation_type,
        )

        if existing_relation is None:
            existing_relation = Relation(
                book_id=book.id,
                source_character_id=source_character.id,
                target_character_id=target_character.id,
                relation_type=relation_type,
                importance_level=payload["importance_level"],
            )
            created_relation_count += 1
        else:
            relation_conflict_count += 1
            if conflict_strategy == WorldConflictStrategy.MANUAL_REVIEW:
                conflicts.append(
                    {
                        "id": f"relation:{existing_relation.id}",
                        "conflict_type": "relation",
                        "target_id": existing_relation.id,
                        "title": f"{source_character.name} -> {target_character.name}",
                        "existing": _serialize_relation_snapshot(existing_relation),
                        "incoming": payload,
                        "source_label": segment.label,
                    }
                )
                continue
            updated_relation_count += 1

        _apply_relation_payload(existing_relation, payload, strategy=conflict_strategy)
        existing_relation.relation_type = relation_type
        if segment.chapter_id is not None:
            if existing_relation.valid_from_chapter_id is None or segment.chapter_id < existing_relation.valid_from_chapter_id:
                existing_relation.valid_from_chapter_id = segment.chapter_id
            if existing_relation.valid_to_chapter_id is None or segment.chapter_id > existing_relation.valid_to_chapter_id:
                existing_relation.valid_to_chapter_id = segment.chapter_id
        db.add(existing_relation)
        db.flush()
        record_relation_event(
            db,
            existing_relation,
            chapter_id=segment.chapter_id,
            segment_label=segment.label,
            relation_type=existing_relation.relation_type,
            label=existing_relation.label,
            description=existing_relation.description,
            strength=existing_relation.strength,
            importance_level=existing_relation.importance_level,
            is_bidirectional=existing_relation.is_bidirectional,
            event_summary=payload["event_summary"] or existing_relation.description,
        )
        touched_relation_ids.add(existing_relation.id)

    existing_factions = db.execute(
        select(Faction).where(Faction.book_id == book.id).order_by(Faction.name.asc(), Faction.id.asc())
    ).scalars().all()
    faction_lookup = {
        ai_service._normalized_name(item.name): item
        for item in existing_factions
        if ai_service._normalized_name(item.name)
    }

    for raw_faction in extracted_payload.factions:
        faction_name = _clean_text(raw_faction.get("name"))
        normalized_faction_name = ai_service._normalized_name(faction_name)
        if not normalized_faction_name:
            continue
        faction = faction_lookup.get(normalized_faction_name)
        if faction is None:
            faction = Faction(
                book_id=book.id,
                name=faction_name,
            )
            created_faction_count += 1
        else:
            updated_faction_count += 1
        if raw_faction.get("description"):
            faction.description = _clean_text(raw_faction.get("description"))
        db.add(faction)
        db.flush()
        touched_faction_ids.add(faction.id)
        faction_lookup[normalized_faction_name] = faction

        raw_memberships = raw_faction.get("memberships") or []
        if not isinstance(raw_memberships, list):
            continue
        for raw_membership in raw_memberships:
            if not isinstance(raw_membership, dict):
                continue
            character_name = ai_service._normalized_name(raw_membership.get("character_name"))
            if not character_name:
                continue
            character = lookup_after_merge.get(character_name)
            if character is None:
                continue
            existing_membership = db.execute(
                select(FactionMembership)
                .where(
                    FactionMembership.book_id == book.id,
                    FactionMembership.faction_id == faction.id,
                    FactionMembership.character_id == character.id,
                )
                .order_by(FactionMembership.id.desc())
            ).scalars().first()
            if existing_membership is None:
                existing_membership = FactionMembership(
                    book_id=book.id,
                    faction_id=faction.id,
                    character_id=character.id,
                    status=normalize_faction_status(raw_membership.get("status")),
                )
            existing_membership.role_label = _clean_text(raw_membership.get("role_label")) or existing_membership.role_label
            existing_membership.loyalty = _safe_parse_strength(raw_membership.get("loyalty"))
            existing_membership.status = normalize_faction_status(raw_membership.get("status"))
            existing_membership.notes = _clean_text(raw_membership.get("notes")) or existing_membership.notes
            if segment.chapter_id is not None and existing_membership.start_chapter_id is None:
                existing_membership.start_chapter_id = segment.chapter_id
            if existing_membership.status == "former" and segment.chapter_id is not None:
                existing_membership.end_chapter_id = segment.chapter_id
            if existing_membership.status == "active":
                existing_membership.end_chapter_id = None
            db.add(existing_membership)
            db.flush()
            touched_membership_ids.add(existing_membership.id)

    appended_world_facts: list[str] = []
    if update_world_bible:
        existing_lines = [line.strip() for line in (book.world_bible or "").splitlines() if line.strip()]
        merged_world_facts, appended_world_facts = ai_service.merge_world_facts(
            existing_lines,
            extracted_payload.world_facts,
        )
        if "\n".join(merged_world_facts) != "\n".join(existing_lines):
            book.world_bible = "\n".join(merged_world_facts)
            db.add(book)

    db.commit()

    return {
        "segment_label": extracted_payload.segment_label,
        "segment_units": extracted_payload.segment_units,
        "characters": merged_character_payloads,
        "relations": merged_relation_payloads,
        "world_facts": extracted_payload.world_facts,
        "world_facts_appended": appended_world_facts,
        "created_character_count": created_character_count,
        "updated_character_count": updated_character_count,
        "created_relation_count": created_relation_count,
        "updated_relation_count": updated_relation_count,
        "created_faction_count": created_faction_count,
        "updated_faction_count": updated_faction_count,
        "character_conflict_count": character_conflict_count,
        "relation_conflict_count": relation_conflict_count,
        "character_ids": sorted(touched_character_ids),
        "relation_ids": sorted(touched_relation_ids),
        "faction_ids": sorted(touched_faction_ids),
        "faction_membership_ids": sorted(touched_membership_ids),
        "conflicts": conflicts,
    }


def run_segment_world_extraction(
    db: Session,
    *,
    book: Book,
    current_user: User,
    segment: ExtractionSegment,
    character_config: ai_service.ResolvedAIConfig,
    relation_config: ai_service.ResolvedAIConfig,
    conflict_strategy: WorldConflictStrategy,
    update_world_bible: bool,
) -> dict[str, Any]:
    existing_characters = db.execute(
        select(Character).where(Character.book_id == book.id).order_by(Character.name.asc())
    ).scalars().all()
    extracted_payload = extract_segment_world_payload(
        book_snapshot=_build_book_prompt_snapshot(book),
        segment=segment,
        existing_character_snapshots=[_serialize_character_prompt_snapshot(item) for item in existing_characters],
        character_config=character_config,
        relation_config=relation_config,
    )
    return apply_segment_world_payload(
        db,
        book=book,
        segment=segment,
        extracted_payload=extracted_payload,
        conflict_strategy=conflict_strategy,
        update_world_bible=update_world_bible,
    )


def _build_internal_blocks(db: Session, job: WorldExtractionJob) -> Iterator[ExtractionBlock]:
    blocks, _ = plan_internal_book_blocks(db, job)
    yield from blocks


def _coalesce_external_blocks(
    source_name: str,
    block_texts: Iterable[str],
    *,
    target_unit_limit: int,
) -> Iterator[ExtractionBlock]:
    normalized_limit = max(MIN_SEGMENT_UNIT_LIMIT, int(target_unit_limit or DEFAULT_SEGMENT_UNIT_LIMIT))
    buffered_parts: list[str] = []
    buffered_units = 0
    start_index: Optional[int] = None
    end_index: Optional[int] = None

    def flush() -> Iterator[ExtractionBlock]:
        nonlocal buffered_parts, buffered_units, start_index, end_index
        if not buffered_parts or start_index is None or end_index is None:
            return
        label = (
            f"{source_name} 第 {start_index} 段"
            if start_index == end_index
            else f"{source_name} 第 {start_index}-{end_index} 段"
        )
        text = "\n\n".join(part for part in buffered_parts if part.strip()).strip()
        if text:
            yield ExtractionBlock(label=label, text=text)
        buffered_parts = []
        buffered_units = 0
        start_index = None
        end_index = None

    for index, raw_text in enumerate(block_texts, start=1):
        text = str(raw_text or "").strip()
        if not text:
            continue
        text_units = ai_service.estimate_text_units(text)
        join_units = 2 if buffered_parts else 0
        if buffered_parts and buffered_units + join_units + text_units > normalized_limit:
            yield from flush()
        if start_index is None:
            start_index = index
        end_index = index
        buffered_parts.append(text)
        buffered_units += join_units + text_units

    yield from flush()


def estimate_import_document(
    path: Path,
    *,
    source_name: Optional[str] = None,
    segment_unit_limit: int,
    update_world_bible: bool,
) -> dict[str, Any]:
    normalized_limit = normalize_segment_unit_limit(segment_unit_limit)
    effective_source_name = source_name or path.name
    raw_block_texts = list(iter_document_blocks(path))
    raw_block_count = len(raw_block_texts)
    raw_text_units = sum(ai_service.estimate_text_units(text) for text in raw_block_texts)
    coalesced_blocks = list(
        _coalesce_external_blocks(
            effective_source_name,
            raw_block_texts,
            target_unit_limit=normalized_limit,
        )
    )
    coalesced_block_count = len(coalesced_blocks)
    estimated_segment_units, estimated_segment_count = count_segments_from_blocks(
        coalesced_blocks,
        segment_unit_limit=normalized_limit,
    )
    estimated_model_calls = estimated_segment_count * 3

    base_input_low = int(estimated_segment_units * IMPORT_ESTIMATE_INPUT_LOW_FACTOR)
    base_input_high = int(estimated_segment_units * IMPORT_ESTIMATE_INPUT_HIGH_FACTOR)
    segment_overhead_low = estimated_segment_count * IMPORT_ESTIMATE_SEGMENT_OVERHEAD_LOW
    segment_overhead_high = estimated_segment_count * IMPORT_ESTIMATE_SEGMENT_OVERHEAD_HIGH
    output_low = max(estimated_segment_count * IMPORT_ESTIMATE_OUTPUT_LOW, 1200 if estimated_segment_count else 0)
    output_high = max(estimated_segment_count * IMPORT_ESTIMATE_OUTPUT_HIGH, 3600 if estimated_segment_count else 0)
    postprocess_low = max(estimated_segment_count * IMPORT_ESTIMATE_POSTPROCESS_LOW, 2000 if estimated_segment_count else 0)
    postprocess_high = max(estimated_segment_count * IMPORT_ESTIMATE_POSTPROCESS_HIGH, 6000 if estimated_segment_count else 0)

    total_low = base_input_low + segment_overhead_low + output_low + postprocess_low
    total_high = base_input_high + segment_overhead_high + output_high + postprocess_high
    if update_world_bible:
        total_low = int(total_low * IMPORT_ESTIMATE_WORLD_BIBLE_LOW_MULTIPLIER)
        total_high = int(total_high * IMPORT_ESTIMATE_WORLD_BIBLE_HIGH_MULTIPLIER)

    assumptions = [
        "估算已按当前导入策略先合并碎段，再按目标长度切片。",
        "每个提取片段默认会触发 3 次模型调用：人物、关系、阵营。",
        "区间已包含提示词、候选人物 JSON、关系/人物后处理摘要等额外开销。",
    ]
    if update_world_bible:
        assumptions.append("已将世界观手册整理与世界事实摘要的额外消耗计入区间。")
    else:
        assumptions.append("未勾选更新世界观手册，因此未计入世界事实整理的额外消耗。")

    return {
        "source_name": effective_source_name,
        "segment_unit_limit": normalized_limit,
        "raw_block_count": raw_block_count,
        "coalesced_block_count": coalesced_block_count,
        "estimated_text_units": estimated_segment_units or raw_text_units,
        "estimated_segment_count": estimated_segment_count,
        "estimated_model_call_count": estimated_model_calls,
        "update_world_bible": update_world_bible,
        "estimated_total_tokens_low": total_low,
        "estimated_total_tokens_high": total_high,
        "assumptions": assumptions,
    }


def _build_external_blocks(job: WorldExtractionJob) -> Iterator[ExtractionBlock]:
    stored_path = Path((job.options_json or {}).get("stored_path") or "")
    if not stored_path.exists():
        raise RuntimeError("The uploaded source document is no longer available.")
    source_name = job.source_name or stored_path.name
    yield from _coalesce_external_blocks(
        source_name,
        iter_document_blocks(stored_path),
        target_unit_limit=job.segment_unit_limit or DEFAULT_SEGMENT_UNIT_LIMIT,
    )


def _job_result_template(job: WorldExtractionJob) -> dict[str, Any]:
    payload = {
        "job_id": job.id,
        "book_id": job.book_id,
        "source_type": job.source_type.value,
        "source_name": job.source_name,
        "conflict_strategy": job_conflict_strategy(job).value,
        "totals": {
            "created_character_count": 0,
            "updated_character_count": 0,
            "created_relation_count": 0,
            "updated_relation_count": 0,
            "created_faction_count": 0,
            "updated_faction_count": 0,
            "world_facts_count": 0,
            "world_facts_appended_count": 0,
            "character_conflict_count": 0,
            "relation_conflict_count": 0,
            "pending_conflict_count": 0,
            "resolved_conflict_count": 0,
            "failed_segment_count": 0,
        },
        "errors": [],
        "items": [],
        "conflicts": [],
        "postprocess": {
            "status": "not_started",
            "characters_summarized": 0,
            "relations_summarized": 0,
            "world_facts_summarized": 0,
        },
    }
    _recount_conflicts(payload)
    return payload


def _chunk_items(items: list[Any], size: int) -> Iterator[list[Any]]:
    chunk_size = max(1, size)
    for index in range(0, len(items), chunk_size):
        yield items[index:index + chunk_size]


def _world_character_summary_prompt(book: Book, characters: list[Character]) -> list[dict[str, str]]:
    payload: list[dict[str, Any]] = []
    for item in characters:
        normalized_card = merge_character_card_json(item.card_json)
        snapshot = {
            "id": item.id,
            "name": item.name,
            "role_label": item.role_label,
            "biography": item.description,
            "traits": item.traits or [],
            "background": item.background,
            "goals": item.goals,
            "secrets": item.secrets,
            "notes": item.notes,
            "life_statuses": normalized_card.get("life_statuses", []),
            "timeline_entries": (normalized_card.get("timeline_entries") or [])[-6:],
        }
        for key in CHARACTER_CARD_TEXT_FIELDS:
            snapshot[key] = normalized_card.get(key)
        payload.append(snapshot)
    return [
        {
            "role": "system",
            "content": (
                "You consolidate extracted novel character cards. Return strict JSON only. "
                "Use concise Chinese when the source is Chinese. "
                "Do not invent facts that are not supported by the provided data."
            ),
        },
        {
            "role": "user",
            "content": (
                "Return a JSON object with this shape:\n"
                "{\n"
                '  "characters": [\n'
                "    {\n"
                '      "id": 1,\n'
                '      "biography": "string"\n'
                "    }\n"
                "  ]\n"
                "}\n\n"
                f"Book title: {book.title}\n"
                f"Character snapshots JSON: {json.dumps(payload, ensure_ascii=False)}\n\n"
                "These snapshots may come from many segment-level partial extractions, so repeated or conflicting details can exist. "
                "Prefer durable facts that repeat or stay consistent, and use the latest timeline state only for current situation. "
                "Rewrite each biography into 1 to 3 coherent Chinese sentences. "
                "Prefer stable identity, relationship position, long-term motive, and current situation. "
                "Use the timeline_entries and life_statuses to keep the biography logically consistent, "
                "but do not copy the timeline as a list or speculate about future plot."
            ),
        },
    ]


def _world_relation_summary_prompt(book: Book, relations: list[Relation]) -> list[dict[str, str]]:
    payload = [
        {
            "id": item.id,
            "source_character_name": item.source_character.name if item.source_character else None,
            "target_character_name": item.target_character.name if item.target_character else None,
            "relation_type": item.relation_type,
            "label": item.label,
            "description": ai_service.normalize_relation_description(item.description),
            "strength": item.strength,
            "importance_level": item.importance_level,
            "is_bidirectional": item.is_bidirectional,
        }
        for item in relations
    ]
    return [
        {
            "role": "system",
            "content": (
                "You consolidate extracted novel relationships. Return strict JSON only. "
                "Keep relation_type unchanged as the provided stable key. "
                "Use Chinese for label and description."
            ),
        },
        {
            "role": "user",
            "content": (
                "Return a JSON object with this shape:\n"
                "{\n"
                '  "relations": [\n'
                "    {\n"
                '      "id": 1,\n'
                '      "label": "string",\n'
                '      "description": "string"\n'
                "    }\n"
                "  ]\n"
                "}\n\n"
                f"Book title: {book.title}\n"
                f"Relation snapshots JSON: {json.dumps(payload, ensure_ascii=False)}\n\n"
                "These relation snapshots may come from many segmented extractions and can contain repeated or noisy wording. "
                "Prefer stable relationship facts, remove weak temporary plot noise, and resolve duplicate wording into one clear statement. "
                "Rewrite only label and description into concise Chinese. "
                "Keep relation_type exactly as given. "
                "Keep them logically consistent and avoid plot recap. "
                "description must stay as one short sentence."
            ),
        },
    ]


def _world_fact_summary_prompt(book: Book, world_facts: list[str]) -> list[dict[str, str]]:
    return [
        {
            "role": "system",
            "content": (
                "You consolidate extracted novel world facts. Return strict JSON only. "
                "Use concise Chinese. Keep only durable setting facts."
            ),
        },
        {
            "role": "user",
            "content": (
                "Return a JSON object with this shape:\n"
                "{\n"
                '  "world_facts": ["string"]\n'
                "}\n\n"
                f"Book title: {book.title}\n"
                f"World facts JSON: {json.dumps(world_facts, ensure_ascii=False)}\n\n"
                "These facts may come from many segmented extractions, so duplicates, overlaps, or local noise may exist. "
                "Prefer durable setting facts that remain valid across the book, and drop one-off plot details unless they define lasting rules. "
                "Rewrite these facts into a cleaner, more logical order from broad setting to specific rules. "
                "Merge duplicates, remove weak temporary details, keep each item as one concise Chinese sentence, "
                "and keep at most 60 items."
            ),
        },
    ]


def _chunk_world_facts_for_summary(world_facts: list[str], *, target_units: int) -> list[list[str]]:
    normalized = ai_service.merge_world_facts([], world_facts)[0]
    chunks: list[list[str]] = []
    current: list[str] = []
    current_units = 0
    limit_units = max(1, target_units)

    for fact in normalized:
        fact_units = max(1, ai_service.estimate_text_units(fact))
        if current and (current_units + fact_units > limit_units or len(current) >= 60):
            chunks.append(current)
            current = []
            current_units = 0
        current.append(fact)
        current_units += fact_units

    if current:
        chunks.append(current)
    return chunks


def _summarize_world_facts_once(
    book: Book,
    *,
    summary_config: ai_service.ResolvedAIConfig,
    world_facts: list[str],
) -> list[str]:
    if not world_facts:
        return []
    payload = _call_strict_json_chat(
        summary_config,
        messages=_world_fact_summary_prompt(book, world_facts),
        expectation="World fact summary response must be a JSON object.",
    )
    if not isinstance(payload, dict):
        return ai_service.merge_world_facts([], world_facts)[0]
    summarized = ai_service.merge_world_facts([], payload.get("world_facts") or [])[0]
    return summarized or ai_service.merge_world_facts([], world_facts)[0]


def _summarize_world_facts_hierarchically(
    book: Book,
    *,
    summary_config: ai_service.ResolvedAIConfig,
    world_facts: list[str],
) -> list[str]:
    current = ai_service.merge_world_facts([], world_facts)[0]
    if not current:
        return []

    passes = 0
    while (
        current
        and (
            len(current) > 60
            or ai_service.estimate_text_units("\n".join(current)) > WORLD_FACT_SUMMARY_MAX_INPUT_UNITS
        )
        and passes < WORLD_FACT_SUMMARY_MAX_PASSES
    ):
        chunks = _chunk_world_facts_for_summary(
            current,
            target_units=WORLD_FACT_SUMMARY_CHUNK_TARGET_UNITS,
        )
        if len(chunks) <= 1:
            break
        condensed: list[str] = []
        for chunk in chunks:
            condensed = ai_service.merge_world_facts(
                condensed,
                _summarize_world_facts_once(
                    book,
                    summary_config=summary_config,
                    world_facts=chunk,
                ),
            )[0]
        if condensed == current:
            break
        current = condensed
        passes += 1

    final_chunks = _chunk_world_facts_for_summary(
        current,
        target_units=WORLD_FACT_SUMMARY_MAX_INPUT_UNITS,
    )
    if len(final_chunks) == 1:
        return _summarize_world_facts_once(
            book,
            summary_config=summary_config,
            world_facts=final_chunks[0],
        )

    merged: list[str] = []
    for chunk in final_chunks:
        merged = ai_service.merge_world_facts(
            merged,
            _summarize_world_facts_once(
                book,
                summary_config=summary_config,
                world_facts=chunk,
            ),
        )[0]

    if merged and (
        len(merged) <= 60
        and ai_service.estimate_text_units("\n".join(merged)) <= WORLD_FACT_SUMMARY_MAX_INPUT_UNITS
    ):
        return _summarize_world_facts_once(
            book,
            summary_config=summary_config,
            world_facts=merged,
        )
    return merged[:60]


def _postprocess_world_extraction_results(
    db: Session,
    *,
    book: Book,
    summary_config: ai_service.ResolvedAIConfig,
    touched_character_ids: set[int],
    touched_relation_ids: set[int],
    update_world_bible: bool,
) -> dict[str, Any]:
    stats = {
        "characters_summarized": 0,
        "relations_summarized": 0,
        "world_facts_summarized": 0,
    }

    if touched_character_ids:
        characters = db.execute(
            select(Character)
            .where(Character.book_id == book.id, Character.id.in_(sorted(touched_character_ids)))
            .order_by(Character.id.asc())
        ).scalars().all()
        for chunk in _chunk_items(characters, 12):
            payload = _call_strict_json_chat(
                summary_config,
                messages=_world_character_summary_prompt(book, chunk),
                expectation="Character summary response must be a JSON object.",
            )
            if not isinstance(payload, dict):
                continue
            updates = payload.get("characters") or []
            updates_by_id = {
                int(item["id"]): item
                for item in updates
                if isinstance(item, dict) and str(item.get("id") or "").strip().isdigit()
            }
            for character in chunk:
                incoming = updates_by_id.get(character.id)
                if not incoming:
                    continue
                biography = _clean_text(incoming.get("biography"))
                if not biography:
                    continue
                character.description = biography
                db.add(character)
                stats["characters_summarized"] += 1

    if touched_relation_ids:
        relations = db.execute(
            select(Relation)
            .where(Relation.book_id == book.id, Relation.id.in_(sorted(touched_relation_ids)))
            .order_by(Relation.id.asc())
        ).scalars().all()
        for chunk in _chunk_items(relations, 20):
            payload = _call_strict_json_chat(
                summary_config,
                messages=_world_relation_summary_prompt(book, chunk),
                expectation="Relation summary response must be a JSON object.",
            )
            if not isinstance(payload, dict):
                continue
            updates = payload.get("relations") or []
            updates_by_id = {
                int(item["id"]): item
                for item in updates
                if isinstance(item, dict) and str(item.get("id") or "").strip().isdigit()
            }
            for relation in chunk:
                incoming = updates_by_id.get(relation.id)
                if not incoming:
                    continue
                incoming_label = _sanitize_relation_text(incoming.get("label"))
                if incoming_label:
                    relation.label = incoming_label
                relation.description = ai_service.normalize_relation_description(incoming.get("description"))
                db.add(relation)
                stats["relations_summarized"] += 1

    if update_world_bible:
        existing_lines = [line.strip() for line in (book.world_bible or "").splitlines() if line.strip()]
        if existing_lines:
            summarized_world_facts = _summarize_world_facts_hierarchically(
                book,
                summary_config=summary_config,
                world_facts=existing_lines,
            )
            if summarized_world_facts:
                book.world_bible = "\n".join(summarized_world_facts)
                db.add(book)
                stats["world_facts_summarized"] = len(summarized_world_facts)

    db.flush()
    return stats


def process_world_extraction_job(job_id: int) -> None:
    db = SessionLocal()
    try:
        job = db.get(WorldExtractionJob, job_id)
        if job is None:
            return
        if job_is_terminated(job):
            return
        if job_cancel_requested(job):
            mark_job_terminated(
                job,
                message="提取任务在开始前已终止。",
            )
            db.add(job)
            db.commit()
            return

        job.status = WorldExtractionJobStatus.RUNNING
        job.started_at = datetime.now(timezone.utc)
        job.message = "正在准备提取计划。"
        _update_job_heartbeat(job, db=db, stage="starting", commit=True)

        book = db.get(Book, job.book_id)
        if book is None:
            raise RuntimeError("The target book no longer exists.")

        actor = job.created_by or book.owner
        if actor is None:
            raise RuntimeError("No valid user context was available for world extraction.")

        job.segment_unit_limit = normalize_segment_unit_limit(job.segment_unit_limit)
        effective_conflict_strategy = job_conflict_strategy(job)
        db.add(job)
        db.commit()

        character_config = ai_service.resolve_ai_config_with_fallback(
            db,
            [
                AIModule.CHARACTER_EXTRACTION,
                AIModule.SETTING_EXTRACTION,
                AIModule.SUMMARY,
                AIModule.CO_WRITING,
            ],
            actor,
            book,
        )
        relation_config = ai_service.resolve_ai_config_with_fallback(
            db,
            [
                AIModule.RELATION_EXTRACTION,
                AIModule.CHARACTER_EXTRACTION,
                AIModule.SETTING_EXTRACTION,
                AIModule.SUMMARY,
                AIModule.CO_WRITING,
            ],
            actor,
            book,
        )
        try:
            summary_config = ai_service.resolve_ai_config_with_fallback(
                db,
                [
                    AIModule.SUMMARY,
                    AIModule.CHARACTER_EXTRACTION,
                    AIModule.SETTING_EXTRACTION,
                    AIModule.RELATION_EXTRACTION,
                    AIModule.CO_WRITING,
                ],
                actor,
                book,
            )
        except ai_service.AIConfigNotFoundError:
            summary_config = character_config

        job.segment_unit_limit, detected_context_window = _context_aware_segment_unit_limit(
            job.segment_unit_limit,
            character_config=character_config,
            relation_config=relation_config,
        )
        options = dict(job.options_json or {})
        options["skip_unchanged_chapters"] = job_skip_unchanged_chapters(job)
        if detected_context_window:
            options["detected_context_window"] = detected_context_window
        job.options_json = options
        _update_job_heartbeat(job, db=db, stage="config_resolved", commit=True)

        chapter_segment_totals: dict[int, int] = {}
        chapter_segment_successes: dict[int, int] = {}
        chapter_signatures: dict[int, str] = {}
        failed_chapter_ids: set[int] = set()
        source_blocks: list[ExtractionBlock]

        if job.source_type == WorldExtractionSource.INTERNAL_BOOK:
            source_blocks, internal_stats = plan_internal_book_blocks(db, job)
            chapter_signatures = dict(internal_stats.get("chapter_signatures") or {})
            options = dict(job.options_json or {})
            options["planned_chapter_count"] = internal_stats["included_chapter_count"]
            options["skipped_empty_chapter_count"] = internal_stats["skipped_empty_chapter_count"]
            options["skipped_unchanged_chapter_count"] = internal_stats["skipped_unchanged_chapter_count"]
            job.options_json = options
            _update_job_heartbeat(job, db=db, stage="planning_internal_blocks", commit=True)
        else:
            source_blocks = list(_build_external_blocks(job))

        planned_segments = build_segment_plan(
            source_blocks,
            segment_unit_limit=job.segment_unit_limit,
        )
        planned_segments = select_job_retry_segments(job, planned_segments)
        chapter_segment_totals = _chapter_segment_totals_from_segments(planned_segments)
        total_units = sum(segment.unit_count for segment in planned_segments)
        total_segments = len(planned_segments)
        worker_count = _effective_worker_count(recommended_worker_count(job), total_segments)

        options = dict(job.options_json or {})
        if job.source_type == WorldExtractionSource.INTERNAL_BOOK:
            options["planned_chapter_count"] = len(
                {segment.chapter_id for segment in planned_segments if segment.chapter_id is not None}
            )
        options["planned_segment_count"] = total_segments
        options["worker_count"] = worker_count
        job.options_json = options
        job.total_units = total_units
        job.total_segments = total_segments
        job.message = (
            "没有找到可重试的失败片段。"
            if (job.options_json or {}).get("retry_failed_only") and total_segments == 0
            else "没有找到可提取的新内容或匹配内容。"
            if total_segments == 0
            else (
                f"已规划 {total_segments} 个提取片段，"
                f"单段约 {job.segment_unit_limit} 字，"
                f"使用 {worker_count} 个线程。"
            )
        )
        _update_job_heartbeat(job, db=db, stage="segment_plan_ready", commit=True)

        result_payload = _job_result_template(job)
        touched_character_ids: set[int] = set()
        touched_relation_ids: set[int] = set()
        if total_segments == 0:
            clear_job_cancel_request(job)
            job.status = WorldExtractionJobStatus.COMPLETED
            job.finished_at = datetime.now(timezone.utc)
            job.result_payload = result_payload
            if (job.options_json or {}).get("retry_failed_only"):
                job.message = "没有找到可重试的失败片段。"
            elif job.source_type == WorldExtractionSource.INTERNAL_BOOK:
                skipped_unchanged = int((job.options_json or {}).get("skipped_unchanged_chapter_count") or 0)
                job.message = (
                    "已写章节都已完成提取，无需重复扫描。"
                    if skipped_unchanged
                    else "没有找到可用于提取的内容。"
                )
            _update_job_heartbeat(job, db=db, stage="completed")
            db.add(job)
            db.commit()
            return

        existing_character_snapshots = [
            _serialize_character_prompt_snapshot(item)
            for item in db.execute(
                select(Character).where(Character.book_id == book.id).order_by(Character.name.asc())
            ).scalars().all()
        ]
        book_snapshot = _build_book_prompt_snapshot(book)
        segment_iterator = iter(planned_segments)

        def commit_processing_heartbeat(segment: ExtractionSegment) -> None:
            db.refresh(job)
            current_index = min(total_segments, max(1, int(job.processed_segments or 0) + 1))
            _update_job_heartbeat(
                job,
                db=db,
                stage="processing",
                message=f"正在处理片段 {current_index}/{job.total_segments}：{segment.label}",
                commit=True,
            )

        def cancellation_requested() -> bool:
            db.refresh(job)
            return job_cancel_requested(job)

        def extract_payload(segment: ExtractionSegment) -> ExtractedSegmentPayload:
            return extract_segment_world_payload(
                book_snapshot=book_snapshot,
                segment=segment,
                existing_character_snapshots=existing_character_snapshots,
                character_config=character_config,
                relation_config=relation_config,
            )

        commit_processing_heartbeat(planned_segments[0])
        for index, (segment, extracted_payload, extraction_error) in enumerate(
            _iter_ordered_parallel_extractions(
                segment_iterator,
                max_workers=worker_count,
                extractor=extract_payload,
                wait_timeout_seconds=float(JOB_HEARTBEAT_INTERVAL_SECONDS),
                on_wait=commit_processing_heartbeat,
                should_stop=cancellation_requested,
            ),
            start=1,
        ):
            db.refresh(job)
            if job_cancel_requested(job):
                mark_job_terminated(
                    job,
                    message="提取任务已按你的要求终止。",
                )
                db.add(job)
                db.commit()
                return
            if extraction_error is not None:
                logger.warning("segment_extraction_failed segment=%s detail=%s", segment.label, extraction_error)
                if segment.chapter_id is not None:
                    failed_chapter_ids.add(segment.chapter_id)
                result_payload["totals"]["failed_segment_count"] += 1
                if len(result_payload["errors"]) < 20:
                    result_payload["errors"].append(
                        {
                            "segment_label": segment.label,
                            "chapter_id": segment.chapter_id,
                            "detail": str(extraction_error),
                        }
                    )
                job.processed_segments = index
                job.processed_units = min(job.total_units, job.processed_units + segment.unit_count)
                job.message = f"已跳过失败片段 {index}/{job.total_segments}：{segment.label}"
                _recount_conflicts(result_payload)
                job.result_payload = result_payload
                _update_job_heartbeat(job, db=db, stage="processing")
                db.add(job)
                db.commit()
                continue
            try:
                segment_result = apply_segment_world_payload(
                    db,
                    book=book,
                    segment=segment,
                    extracted_payload=extracted_payload,
                    conflict_strategy=effective_conflict_strategy,
                    update_world_bible=job.update_world_bible,
                )
                result_payload["items"].append(
                    {
                        "segment_label": segment_result["segment_label"],
                        "segment_units": segment_result["segment_units"],
                        "created_character_count": segment_result["created_character_count"],
                        "updated_character_count": segment_result["updated_character_count"],
                        "created_relation_count": segment_result["created_relation_count"],
                        "updated_relation_count": segment_result["updated_relation_count"],
                        "created_faction_count": segment_result.get("created_faction_count", 0),
                        "updated_faction_count": segment_result.get("updated_faction_count", 0),
                        "world_facts_count": len(segment_result["world_facts"]),
                        "world_facts_appended_count": len(segment_result["world_facts_appended"]),
                        "character_conflict_count": segment_result["character_conflict_count"],
                        "relation_conflict_count": segment_result["relation_conflict_count"],
                    }
                )
                totals = result_payload["totals"]
                totals["created_character_count"] += segment_result["created_character_count"]
                totals["updated_character_count"] += segment_result["updated_character_count"]
                totals["created_relation_count"] += segment_result["created_relation_count"]
                totals["updated_relation_count"] += segment_result["updated_relation_count"]
                totals["created_faction_count"] += segment_result.get("created_faction_count", 0)
                totals["updated_faction_count"] += segment_result.get("updated_faction_count", 0)
                totals["world_facts_count"] += len(segment_result["world_facts"])
                totals["world_facts_appended_count"] += len(segment_result["world_facts_appended"])
                totals["character_conflict_count"] += segment_result["character_conflict_count"]
                totals["relation_conflict_count"] += segment_result["relation_conflict_count"]
                touched_character_ids.update(
                    int(item)
                    for item in (segment_result.get("character_ids") or [])
                    if str(item).isdigit()
                )
                touched_relation_ids.update(
                    int(item)
                    for item in (segment_result.get("relation_ids") or [])
                    if str(item).isdigit()
                )
                for conflict in segment_result.get("conflicts", []):
                    _add_or_update_conflict(result_payload["conflicts"], conflict)
                if (
                    segment.chapter_id is not None
                    and segment.chapter_id not in failed_chapter_ids
                    and segment.chapter_id in chapter_segment_totals
                ):
                    chapter_segment_successes[segment.chapter_id] = (
                        chapter_segment_successes.get(segment.chapter_id, 0) + 1
                    )
                    if chapter_segment_successes[segment.chapter_id] >= chapter_segment_totals[segment.chapter_id]:
                        chapter = db.get(Chapter, segment.chapter_id)
                        if chapter is not None:
                            mark_chapter_internal_world_extracted(
                                chapter,
                                job_id=job.id,
                                signature=chapter_signatures.get(segment.chapter_id),
                            )
                            db.add(chapter)
            except (ai_service.AIInvocationError, ValueError, KeyError, TypeError) as exc:
                logger.warning(
                    "segment_business_error segment=%s error_type=%s detail=%s",
                    segment.label,
                    type(exc).__name__,
                    exc,
                )
                db.rollback()
                if segment.chapter_id is not None:
                    failed_chapter_ids.add(segment.chapter_id)
                result_payload["totals"]["failed_segment_count"] += 1
                if len(result_payload["errors"]) < 20:
                    result_payload["errors"].append(
                        {
                            "segment_label": segment.label,
                            "chapter_id": segment.chapter_id,
                            "detail": str(exc),
                        }
                    )
            except (RuntimeError, OSError):
                db.rollback()
                logger.exception("segment_system_error segment=%s", segment.label)
                raise
            except Exception:
                db.rollback()
                logger.exception("segment_unexpected_error segment=%s", segment.label)
                raise

            job.processed_segments = index
            job.processed_units = min(job.total_units, job.processed_units + segment.unit_count)
            job.message = f"已处理片段 {index}/{job.total_segments}：{segment.label}"
            _recount_conflicts(result_payload)
            job.result_payload = result_payload
            _update_job_heartbeat(job, db=db, stage="processing")
            db.add(job)
            db.commit()

        db.refresh(job)
        if touched_character_ids or touched_relation_ids or job.update_world_bible:
            _update_job_heartbeat(
                job,
                db=db,
                stage="postprocessing",
                message="正在整理提取结果并总结人物关系。",
                commit=True,
            )
            try:
                result_payload["postprocess"] = {
                    "status": "completed",
                    **_postprocess_world_extraction_results(
                        db,
                        book=book,
                        summary_config=summary_config,
                        touched_character_ids=touched_character_ids,
                        touched_relation_ids=touched_relation_ids,
                        update_world_bible=job.update_world_bible,
                    ),
                }
            except Exception as exc:
                db.rollback()
                logger.exception("world_extraction_postprocess_failed job_id=%s", job.id)
                result_payload["postprocess"] = {
                    "status": "failed",
                    "characters_summarized": 0,
                    "relations_summarized": 0,
                    "world_facts_summarized": 0,
                    "error": str(exc),
                }
                book = db.get(Book, job.book_id) or book
                job = db.get(WorldExtractionJob, job.id) or job
        else:
            result_payload["postprocess"]["status"] = "skipped"

        db.refresh(job)
        clear_job_cancel_request(job)
        job.status = WorldExtractionJobStatus.COMPLETED
        job.finished_at = datetime.now(timezone.utc)
        _recount_conflicts(result_payload)
        postprocess_status = (result_payload.get("postprocess") or {}).get("status")
        postprocess_counts = result_payload.get("postprocess") or {}
        postprocess_applied = any(
            int(postprocess_counts.get(key) or 0) > 0
            for key in ("characters_summarized", "relations_summarized", "world_facts_summarized")
        )
        job.message = (
            f"提取已完成，其中 {result_payload['totals']['failed_segment_count']} 个片段失败。"
            if result_payload["totals"]["failed_segment_count"]
            else "提取任务已完成，并已整理结果摘要。"
            if postprocess_status == "completed" and postprocess_applied
            else "提取任务已完成，但摘要整理失败。"
            if postprocess_status == "failed"
            else "提取任务已完成。"
        )
        job.result_payload = result_payload
        _update_job_heartbeat(job, db=db, stage="completed")
        db.add(job)
        db.commit()
    except WorldExtractionCancellationRequested:
        db.rollback()
        cancelled_job = db.get(WorldExtractionJob, job_id)
        if cancelled_job is not None:
            mark_job_terminated(
                cancelled_job,
                message="提取任务已按你的要求终止。",
            )
            db.add(cancelled_job)
            db.commit()
    except Exception as exc:
        db.rollback()
        failed_job = db.get(WorldExtractionJob, job_id)
        if failed_job is not None:
            clear_job_cancel_request(failed_job)
            failed_job.status = WorldExtractionJobStatus.FAILED
            failed_job.finished_at = datetime.now(timezone.utc)
            failed_job.error_message = str(exc)
            failed_job.message = "提取任务失败。"
            _update_job_heartbeat(failed_job, db=db, stage="failed")
            db.add(failed_job)
            db.commit()
    finally:
        db.close()


def resolve_world_extraction_conflict(
    db: Session,
    *,
    job: WorldExtractionJob,
    conflict_id: str,
    decision: WorldConflictStrategy,
) -> dict[str, Any]:
    if decision not in {WorldConflictStrategy.KEEP_EXISTING, WorldConflictStrategy.PREFER_IMPORTED}:
        raise RuntimeError("Only keep_existing or prefer_imported can be used to resolve a conflict.")

    result_payload = copy.deepcopy(job.result_payload or {})
    conflicts = list(result_payload.get("conflicts") or [])
    target = next((item for item in conflicts if item.get("id") == conflict_id), None)
    if target is None:
        raise RuntimeError("Conflict record was not found.")
    if target.get("status") == "resolved":
        return target

    if decision == WorldConflictStrategy.PREFER_IMPORTED:
        if target.get("conflict_type") == "character":
            character = db.get(Character, target.get("target_id"))
            if character is None:
                raise RuntimeError("The existing character could not be found.")
            incoming = target.get("incoming") or {}
            character.aliases = incoming.get("aliases") or []
            character.role_label = incoming.get("role_label")
            character.description = incoming.get("description")
            character.traits = incoming.get("traits") or []
            character.background = incoming.get("background")
            character.goals = incoming.get("goals")
            character.secrets = incoming.get("secrets")
            character.notes = incoming.get("notes")
            character.first_appearance_chapter_id = incoming.get("first_appearance_chapter_id")
            character.last_appearance_chapter_id = incoming.get("last_appearance_chapter_id")
            character.is_active = True
            character.card_json = incoming.get("card_json") or {}
            db.add(character)
        elif target.get("conflict_type") == "relation":
            relation = db.get(Relation, target.get("target_id"))
            if relation is None:
                raise RuntimeError("The existing relation could not be found.")
            incoming = target.get("incoming") or {}
            relation.label = incoming.get("label")
            relation.description = ai_service.normalize_relation_description(incoming.get("description"))
            relation.strength = _safe_parse_strength(incoming.get("strength"))
            relation.importance_level = normalize_relation_importance(incoming.get("importance_level"))
            relation.is_bidirectional = _parse_boolean(incoming.get("is_bidirectional"))
            db.add(relation)
            db.flush()
            record_relation_event(
                db,
                relation,
                chapter_id=relation.valid_to_chapter_id,
                segment_label="conflict_resolution",
                relation_type=relation.relation_type,
                label=relation.label,
                description=relation.description,
                strength=relation.strength,
                importance_level=relation.importance_level,
                is_bidirectional=relation.is_bidirectional,
                event_summary=ai_service.normalize_relation_description(incoming.get("event_summary") or relation.description),
            )
        else:
            raise RuntimeError("Unsupported conflict type.")

    target["status"] = "resolved"
    target["resolution"] = decision.value
    target["resolved_at"] = _now_iso()
    target["updated_at"] = target["resolved_at"]
    result_payload["conflicts"] = conflicts
    _recount_conflicts(result_payload)
    job.result_payload = result_payload
    job.message = f"已处理冲突 {conflict_id}。"
    db.add(job)
    db.commit()
    db.refresh(job)
    return target
